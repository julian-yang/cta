from docx import Document
from docx.shared import Inches
import lib.book_pb2 as book_pb2
import my_firebase as firebase
import utils
from article_utils import parse_firebase_article
from scripts.query_articles import where_tag


def write_book_to_docx(book):
    document = Document()
    document.add_heading(book.title, 0)
    document.add_page_break()
    for chapter in book.chapters:
        document.add_heading(chapter.title)
        for paragraph in chapter.paragraphs:
            document.add_paragraph(paragraph)
        document.add_page_break()

    filename = f'{book.title}.docx'
    print(f'Writing docx file: {filename}')
    document.save(filename)


def create_chapter_proto(article):
    chapter = book_pb2.Chapter()
    chapter.title = article.chinese_title
    chapter.paragraphs.extend(article.chinese_body.split('\n'))
    return chapter


def create_book_proto(book_title, book_articles):
    book = book_pb2.Book()
    book.title = book_title
    book_articles.sort(key=lambda article: article.chapter_num)
    book.chapters.extend([create_chapter_proto(article) for article in book_articles])
    return book


def get_book_articles_for_tag(db, book_tag):
    docs = firebase.get_existing_articles(db, [where_tag(book_tag)])
    book_articles = [parse_firebase_article(doc) for doc in docs]
    return book_articles


def get_possible_book_tags(db):
    docs = firebase.get_existing_articles(db, [where_tag(u'book')])
    book_articles = [parse_firebase_article(doc) for doc in docs]
    tags = list(set([tag for article in book_articles for tag in article.tags if tag not in ['book', 'aixdzs', 'dushu']]))
    return tags

def print_possible_book_tags(tags):
    print('Possible book tags: ')
    for i in range(len(tags)):
        print(f' * ({i+1}/{len(tags)}) {tags[i]}')


if __name__ == "__main__":
    db = firebase.get_db()

    book_articles = []
    book_tag = ''
    while not book_articles:
        possible_tags = get_possible_book_tags(db)
        print_possible_book_tags(possible_tags)
        index = int(input(f'Enter book tag (#): ')) - 1
        book_tag = possible_tags[index]
        utils.get_yes_no(f'Confirm book tag (y/n): {book_tag}\n')
        book_articles = get_book_articles_for_tag(db, book_tag)
        if not book_articles:
            print(f'Could not find any scraped_articles with tag: {book_tag}!')

    book = create_book_proto(book_tag, book_articles)
    write_book_to_docx(book)
    print('done!')

